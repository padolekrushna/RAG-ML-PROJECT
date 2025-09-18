import streamlit as st
import PyPDF2
import docx
from typing import List, IO
import re

class DocumentProcessor:
    """Handles processing of different document types"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_document(self, file) -> List[str]:
        """
        Process uploaded document and return text chunks
        
        Args:
            file: Streamlit uploaded file object
            
        Returns:
            List of text chunks
        """
        try:
            # Determine file type and extract text
            if file.type == "application/pdf":
                text = self._extract_pdf_text(file)
            elif file.type == "text/plain":
                text = self._extract_txt_text(file)
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = self._extract_docx_text(file)
            else:
                raise ValueError(f"Unsupported file type: {file.type}")
            
            # Clean and chunk the text
            cleaned_text = self._clean_text(text)
            chunks = self._create_chunks(cleaned_text, file.name)
            
            return chunks
            
        except Exception as e:
            st.error(f"Error processing {file.name}: {str(e)}")
            return []
    
    def _extract_pdf_text(self, file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_txt_text(self, file) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file.seek(0)  # Reset file pointer
                    text = file.read().decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def _extract_docx_text(self, file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-"]', ' ', text)
        
        # Remove extra spaces
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def _create_chunks(self, text: str, source_name: str) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Input text
            source_name: Name of source document
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find end position
            end = start + self.chunk_size
            
            # If this isn't the last chunk, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence end
                sentence_end = text.rfind('.', start, end)
                if sentence_end != -1 and sentence_end > start + self.chunk_size // 2:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    word_end = text.rfind(' ', start, end)
                    if word_end != -1 and word_end > start + self.chunk_size // 2:
                        end = word_end
            
            # Extract chunk
            chunk = text[start:end].strip()
            
            if chunk:
                # Add metadata to chunk
                chunk_with_metadata = f"[Source: {source_name}]\n{chunk}"
                chunks.append(chunk_with_metadata)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            
            # Ensure we don't go backwards
            if start <= 0:
                start = end
        
        return chunks
    
    def get_chunk_info(self, chunks: List[str]) -> dict:
        """Get information about the chunks"""
        if not chunks:
            return {
                'total_chunks': 0,
                'avg_chunk_length': 0,
                'total_characters': 0,
                'sources': []
            }
        
        total_chars = sum(len(chunk) for chunk in chunks)
        avg_length = total_chars // len(chunks) if chunks else 0
        
        # Extract source names
        sources = set()
        for chunk in chunks:
            if chunk.startswith('[Source:'):
                source_line = chunk.split('\n')[0]
                source_name = source_line.replace('[Source:', '').replace(']', '').strip()
                sources.add(source_name)
        
        return {
            'total_chunks': len(chunks),
            'avg_chunk_length': avg_length,
            'total_characters': total_chars,
            'sources': list(sources)
        }